#!/usr/bin/env python3
"""Synchronize the two-page DOCX brief with the public story copy."""
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "EvoForecast-decision-brief.docx"

REPLACEMENTS = {
    0: "GRAPH OF LIFE   |   EVOFORECAST                                        DECISION BRIEF · 21 JUL 2026",
    6: "A programme can buy more sequencing, more time points, and more replicates while leaving two forecasting designs indistinguishable. Averages hide the failure that matters. A model can post a strong mean error and still miss often enough to fail at the moment of decision. This artifact uses tens of millions of US dollars as illustrative programme scale. At that scale, waiting for field data can lock the design before the comparison is resolved. This design problem is cheapest to solve before anyone collects a sample.",
    8: "The engine builds a simulated world where the true evolutionary outcome is fixed by construction, then seals it. Competing measurement designs receive registered data packages and modeled-resource allowances. They deposit forecasts before reveal. The harness then scores every deposit on the same footing, covering accuracy, calibration of stated uncertainty, and whether design rankings survive perturbation.",
    9: "It returns three answers that become irrecoverable once collection begins. It shows whether rival designs can be told apart at an affordable sample size. It reveals which added measurement makes those designs distinguishable, a property called identifiability. It also returns a redesign decision when the evidence remains weak. A round that ends in redesign has done its job.",
    11: "The frozen round completed 5,600 of 5,600 simulated runs successfully. The genome-plus-phenome track, which combines genetic information with measured traits, improved mean accuracy by 47.5 percent over the pedigree baseline, aggregated across 35 contexts. Its nominal 90 percent interval covered 80 percent of held-out values. The feedback track added 0.0 percent beyond genome plus phenome. Zero of 60 forward-looking cells reached the 50 percent combined gate. One of 80 cells reached it, a day-90 nowcast at 52.5 percent. Zero of 80 were robust. Zero of 120 experiment plans were robustly best. The favorable control returned +17.0 percent and the matched adverse control returned -9.2 percent. Both are marker-encoded software controls.",
    12: "The clean run gives a clear instruction. Delay scale-up and design a stronger test.",
    14: "Every displayed statistic points to a frozen source row and a recorded SHA-256 fingerprint. The browser is a lookup surface. Forecasting, model fitting, simulation, scoring, and access to latent truth stay upstream. The favorable and adverse controls move in opposite directions, showing that the workflow can expose failure.",
    15: "SLiM 5.2 is the registered trajectory engine for this round. The surrounding harness is engine-agnostic, and other evolutionary simulators are under evaluation. Ecological network work in a population context is a development direction for later rounds.",
    17: "The wind tunnel is an early design stage ahead of field seasons, facilities, and staff. Before several teams commit to a shared measurement standard, it can show whether that standard would let them distinguish their forecasts.",
    18: "A new interactive view exposes all 32 registered trajectories in one Study E persistence cell. Each vessel starts with the same 80 grazers and eight immutable founders. All retain grazers through day 90, while standing selection produces different population and founder histories. That distribution makes the forecasting problem tangible. The broader instrument then compares modeled resource use with forecast value and shows when a data layer changes accuracy, identifiability, or both.",
    20: "One candidate keystone organism. One contained stressor. One spendable budget in US dollars. One outcome sealed until scoring. Rehearse the design, then run the smallest real round that can estimate calibration and forecast accuracy under an independent reveal. A later contained round can test whether a targeted intervention improves resilience. Interacting species, spatial structure, small model ecosystems, and molecular layers expand after the forecast layer earns trust.",
}


def replace_text(paragraph, value: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(value)
        return
    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_cell(cell, value: str) -> None:
    replace_text(cell.paragraphs[0], value)
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def main() -> int:
    document = Document(PATH)
    if len(document.paragraphs) != 28 or len(document.tables) not in (2, 3):
        raise SystemExit("unexpected decision-brief structure")
    for index, value in REPLACEMENTS.items():
        replace_text(document.paragraphs[index], value)
    replace_cell(document.tables[0].cell(0, 0), "5,600 / 5,600\nsimulated runs completed successfully")
    replace_cell(document.tables[1].cell(0, 0), "Claim boundary. All evidence is synthetic and within-model. Real-organism forecast skill, real DNA predictive value, out-of-distribution transfer, and intervention efficacy await empirical testing. Independent review precedes any operational, partnership, or funding claim.")
    author_table = document.tables[2] if len(document.tables) == 3 else document.add_table(rows=1, cols=2)
    replace_cell(author_table.cell(0, 0), "Jay Gutierrez, PhD\nBuilt through GraphOfLife across ecological network science, graph intelligence, evolutionary forecasting, nature-finance translation, and structural-risk publishing.")
    replace_cell(author_table.cell(0, 1), "CONTACT\njg@graphoflife.com\nbiome-translator.emergent.host")
    document.core_properties.modified = document.core_properties.created
    document.save(PATH)
    print("updated", PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
